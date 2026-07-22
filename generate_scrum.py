from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_scrum_document():
    doc = Document()
    
    # Title Page
    doc.add_heading('PHASE 2: SCRUM ARTIFACTS', 0)
    title = doc.add_paragraph('Project Title: Hire Me - A Local Service Marketplace\n')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # SECTION 1: PRODUCT BACKLOG
    doc.add_heading('1. Product Backlog', level=1)
    doc.add_paragraph("The Product Backlog is an ordered, prioritized list of all the features, functions, and requirements necessary to build the 'Hire Me' platform.")
    
    # Table for Product Backlog
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'User Story / Feature'
    hdr_cells[2].text = 'Priority'
    hdr_cells[3].text = 'Estimate (Story Points)'
    
    backlog_items = [
        ("PB-01", "As a User, I want to securely register and log in so that my data is protected.", "High", "3"),
        ("PB-02", "As an Admin, I want to approve/reject Provider verification documents so that only trusted professionals are listed.", "High", "5"),
        ("PB-03", "As a Client, I want to search and filter providers by category and price so that I can find the right fit.", "Medium", "3"),
        ("PB-04", "As a User, I want to top-up a digital wallet via a Mobile Money simulation so that I can pay for services.", "High", "5"),
        ("PB-05", "As a Client, I want my funds held in an Escrow system when I book a job so that I am protected from fraud.", "High", "8"),
        ("PB-06", "As a Provider, I want funds released automatically to my wallet when a job is marked completed.", "High", "8"),
        ("PB-07", "As a Client/Provider, I want to chat in real-time so that we can coordinate the physical job details.", "Medium", "5"),
        ("PB-08", "As an Admin, I want to view platform analytics (total users, total bookings) to monitor growth.", "Low", "2")
    ]
    
    for item in backlog_items:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
        row_cells[3].text = item[3]
        
    doc.add_page_break()

    # SECTION 2: SPRINT BACKLOG
    doc.add_heading('2. Sprint Backlog', level=1)
    doc.add_paragraph("The Sprint Backlog breaks down the prioritized Product Backlog features into granular, actionable engineering tasks distributed across two-week iterative sprints.")

    # Sprint 1
    doc.add_heading('Sprint 1: Core Foundation & Authentication', level=2)
    doc.add_paragraph("Goal: Establish the database architecture, initialize the repositories, and build the secure JWT authentication flow.")
    s1_table = doc.add_table(rows=1, cols=3)
    s1_table.style = 'Light Shading Accent 1'
    s1_hdr = s1_table.rows[0].cells
    s1_hdr[0].text = 'Task ID'
    s1_hdr[1].text = 'Actionable Task Description'
    s1_hdr[2].text = 'Status'
    
    s1_tasks = [
        ("T-1.1", "Initialize Node.js/Express backend and React.js frontend structures.", "Done"),
        ("T-1.2", "Design relational schema for User and Booking models using Sequelize ORM.", "Done"),
        ("T-1.3", "Implement bcrypt password hashing and JWT token generation for Login API.", "Done"),
        ("T-1.4", "Build responsive Landing/Auth frontend page with Tailwind CSS.", "Done")
    ]
    for t in s1_tasks:
        row = s1_table.add_row().cells
        row[0].text = t[0]
        row[1].text = t[1]
        row[2].text = t[2]

    doc.add_paragraph("") # Spacing

    # Sprint 2
    doc.add_heading('Sprint 2: Profiles & Provider Verification', level=2)
    doc.add_paragraph("Goal: Enable service providers to build their profiles and allow administrators to vet and verify their identity documents.")
    s2_table = doc.add_table(rows=1, cols=3)
    s2_table.style = 'Light Shading Accent 2'
    s2_hdr = s2_table.rows[0].cells
    s2_hdr[0].text = 'Task ID'
    s2_hdr[1].text = 'Actionable Task Description'
    s2_hdr[2].text = 'Status'
    
    s2_tasks = [
        ("T-2.1", "Create ProviderProfile model with a One-to-One relationship to User.", "Done"),
        ("T-2.2", "Configure multer middleware for handling multipart/form-data document uploads.", "Done"),
        ("T-2.3", "Develop the Admin Dashboard UI to list 'pending verification' providers.", "Done"),
        ("T-2.4", "Implement PATCH /api/admin/providers/:id/verify endpoint.", "Done")
    ]
    for t in s2_tasks:
        row = s2_table.add_row().cells
        row[0].text = t[0]
        row[1].text = t[1]
        row[2].text = t[2]

    doc.add_paragraph("")

    # Sprint 3
    doc.add_heading('Sprint 3: Wallet System & Escrow Booking', level=2)
    doc.add_paragraph("Goal: Implement the financial core of the platform, ensuring transactional atomicity and fraud protection via Escrow holds.")
    s3_table = doc.add_table(rows=1, cols=3)
    s3_table.style = 'Light Shading Accent 3'
    s3_hdr = s3_table.rows[0].cells
    s3_hdr[0].text = 'Task ID'
    s3_hdr[1].text = 'Actionable Task Description'
    s3_hdr[2].text = 'Status'
    
    s3_tasks = [
        ("T-3.1", "Add 'wallet_balance' column and create the Transaction audit ledger model.", "Done"),
        ("T-3.2", "Build Wallet Top-Up UI simulating 9-digit Cameroonian Mobile Money.", "Done"),
        ("T-3.3", "Write DB Transactions to deduct funds instantly when a Booking is created.", "Done"),
        ("T-3.4", "Write completion logic to atomically release funds to Provider's wallet.", "Done")
    ]
    for t in s3_tasks:
        row = s3_table.add_row().cells
        row[0].text = t[0]
        row[1].text = t[1]
        row[2].text = t[2]

    doc.add_paragraph("")

    # Sprint 4
    doc.add_heading('Sprint 4: Search & Real-Time Chat', level=2)
    doc.add_paragraph("Goal: Build the Client discovery interface and integrate WebSockets for instant communication between matched users.")
    s4_table = doc.add_table(rows=1, cols=3)
    s4_table.style = 'Light Shading Accent 4'
    s4_hdr = s4_table.rows[0].cells
    s4_hdr[0].text = 'Task ID'
    s4_hdr[1].text = 'Actionable Task Description'
    s4_hdr[2].text = 'Status'
    
    s4_tasks = [
        ("T-4.1", "Build Client Dashboard React component with category/price filters.", "Done"),
        ("T-4.2", "Implement backend GET /api/providers with query filtering logic.", "Done"),
        ("T-4.3", "Set up Socket.io server and configure secure JWT WebSocket authentication.", "Done"),
        ("T-4.4", "Build floating Chat Modal UI to display real-time message streams.", "Done")
    ]
    for t in s4_tasks:
        row = s4_table.add_row().cells
        row[0].text = t[0]
        row[1].text = t[1]
        row[2].text = t[2]

    doc.save('Hire_Me_Scrum_Artifacts.docx')
    print("Scrum Artifacts Document generated successfully: Hire_Me_Scrum_Artifacts.docx")

if __name__ == '__main__':
    create_scrum_document()
