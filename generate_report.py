from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_report():
    doc = Document()
    
    # Title Page
    doc.add_heading('COMPREHENSIVE PROJECT REPORT', 0)
    title = doc.add_paragraph('Project Title: Hire Me - A Local Service Marketplace\n')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # CHAPTER ONE
    doc.add_heading('Chapter One: Introduction', level=1)
    
    doc.add_heading('1.1 General Introduction', level=2)
    doc.add_paragraph(
        "The rapid growth of the gig economy has fundamentally changed how individuals access and offer services. "
        "However, finding reliable, verified professionals for everyday physical tasks (such as plumbing, cleaning, or electrical work) "
        "remains a significant challenge in many local communities. Traditional word-of-mouth methods are often inefficient, "
        "and existing digital platforms frequently lack robust verification and secure localized payment systems. "
        "The 'Hire Me' platform was conceived to bridge this gap, offering a dedicated marketplace that connects clients "
        "with skilled, vetted professionals while ensuring secure and transparent transactions."
    )

    doc.add_heading('1.2 Aim and Objectives of "Hire Me"', level=2)
    doc.add_paragraph("The primary aim of the 'Hire Me' project is to design and develop a full-stack web application that facilitates the seamless discovery, booking, and payment of local service professionals.")
    doc.add_paragraph("The specific objectives include:", style='List Bullet')
    doc.add_paragraph("To develop a secure authentication and role-based access system (Admin, Client, Provider).", style='List Bullet')
    doc.add_paragraph("To implement an Escrow-based wallet system that protects both the client's funds and the provider's earnings until the job is satisfactorily completed.", style='List Bullet')
    doc.add_paragraph("To integrate a real-time chat infrastructure enabling instant communication between clients and providers.", style='List Bullet')
    doc.add_paragraph("To design a highly responsive, mobile-first user interface for accessible usage across all devices.", style='List Bullet')

    doc.add_heading('1.3 Problem Statement', level=2)
    doc.add_paragraph(
        "Currently, clients face significant risks when hiring local professionals due to a lack of identity verification "
        "and quality assurance. Simultaneously, service providers struggle with payment defaults and inconsistent client "
        "communication. There is a lack of a centralized, secure platform that inherently enforces trust through a digital "
        "escrow system and transparent review mechanisms. 'Hire Me' directly addresses these vulnerabilities by mandating "
        "provider documentation verification and holding funds securely during active jobs."
    )
    doc.add_page_break()

    # CHAPTER TWO
    doc.add_heading('Chapter Two: Literature Review', level=1)
    
    doc.add_heading('2.1 Review of Software Development Methodologies', level=2)
    doc.add_paragraph(
        "Several methodologies exist for software engineering, prominently Waterfall, Kanban, and Agile (Scrum). "
        "The Waterfall model is highly structured and sequential, making it unsuitable for projects where requirements "
        "are expected to evolve. Kanban focuses on continuous delivery without fixed iterations. Agile frameworks, "
        "particularly Scrum, emphasize iterative development, regular feedback loops, and cross-functional collaboration."
    )

    doc.add_heading('2.2 Comparison and Justification for Scrum', level=2)
    doc.add_paragraph(
        "Unlike Waterfall, which delays testing until the final phases, Scrum allows for continuous integration and testing. "
        "For 'Hire Me', Scrum was chosen as the optimal methodology because the platform required rapid prototyping and "
        "frequent adjustments based on user interface testing and database schema evolution. The iterative nature of Scrum "
        "allowed the team to deliver a functional Minimum Viable Product (MVP) quickly, such as basic authentication, "
        "before incrementally adding complex features like the Escrow wallet and real-time WebSockets."
    )

    doc.add_heading('2.3 Review of Related Literature', level=2)
    doc.add_paragraph(
        "Existing literature and market analysis reveal that while global freelance platforms (e.g., Upwork, Fiverr) excel "
        "in digital service delivery, they are ill-equipped for local, physical services. Platforms like TaskRabbit exist "
        "in Western markets but often lack integration with localized payment methods (e.g., Mobile Money simulators) and "
        "rigorous decentralized provider verification. 'Hire Me' innovates on these models by combining localized financial "
        "workflows with an automated administrative approval matrix."
    )
    doc.add_page_break()

    # CHAPTER THREE
    doc.add_heading('Chapter Three: Methodology and Materials', level=1)

    doc.add_heading('3.1 System Requirements', level=2)
    doc.add_heading('Functional Requirements', level=3)
    doc.add_paragraph("1. The system must allow users to register as either Clients or Providers.")
    doc.add_paragraph("2. Providers must be able to upload verification documents for Admin approval.")
    doc.add_paragraph("3. The system must feature a digital wallet enabling users to top-up funds.")
    doc.add_paragraph("4. Bookings must lock client funds in Escrow and release them only upon job completion.")
    doc.add_paragraph("5. Clients and Providers must be able to communicate via real-time chat.")

    doc.add_heading('Non-Functional Requirements', level=3)
    doc.add_paragraph("- Security: Passwords must be hashed using bcrypt, and API routes secured via JWT.")
    doc.add_paragraph("- Performance: Real-time chat messages must be delivered with sub-second latency via Socket.io.")
    doc.add_paragraph("- Usability: The application must be fully responsive, supporting mobile and desktop views.")

    doc.add_heading('3.2 High-Level System Architecture (HLD)', level=2)
    doc.add_paragraph(
        "The system employs a standard 3-tier Client-Server architecture. The Presentation Layer is built as a Single Page "
        "Application (SPA) using React. The Application Layer consists of a Node.js/Express REST API. The Data Layer is "
        "managed by a MySQL relational database interacted with via the Sequelize ORM."
    )
    
    doc.add_heading('3.3 UML Diagrams', level=2)
    doc.add_paragraph("[INSERT USE CASE DIAGRAM HERE]")
    doc.add_paragraph("Figure 3.1: Use Case Diagram illustrating interactions between Clients, Providers, and Admins.\n")
    doc.add_paragraph("[INSERT CLASS DIAGRAM HERE]")
    doc.add_paragraph("Figure 3.2: Class Diagram showing User, ProviderProfile, Booking, Message, and Transaction models.\n")
    doc.add_paragraph("[INSERT SEQUENCE DIAGRAM HERE]")
    doc.add_paragraph("Figure 3.3: Sequence Diagram illustrating the Escrow payment flow during a booking.\n")

    doc.add_heading('3.4 Application of Scrum', level=2)
    doc.add_heading('Team Organization', level=3)
    doc.add_paragraph(
        "The project was executed by a cross-functional team of 4 members, organized as follows:\n"
        "1. Scrum Master and Backend Lead: Managed sprint planning, unblocked technical hurdles, and architected the API.\n"
        "2. Frontend Lead and DevOps Engineer: Built the responsive UI and configured deployment pipelines.\n"
        "3. System Analyst and Lead Technical Writer: Gathered requirements, designed the database schema, and authored documentation.\n"
        "4. UML Architect and Quality Assurance Engineer: Designed system models and conducted end-to-end testing."
    )
    
    doc.add_heading('Workflow Management', level=3)
    doc.add_paragraph("Work was divided into two-week sprints. We utilized daily asynchronous stand-ups to track progress and adapted a kanban board for task progression (To Do, In Progress, Review, Done).")

    doc.add_heading('Conflict Resolution and Challenges Faced', level=3)
    doc.add_paragraph(
        "The team faced several significant challenges, primarily revolving around technical complexities and time constraints. "
        "A major coding challenge occurred during the database migration from SQLite to MySQL, which caused schema synchronization conflicts. "
        "Additionally, integrating the Wallet Escrow logic with asynchronous booking states led to race conditions. "
        "To resolve these conflicts, the Backend Lead and QA Engineer pair-programmed to implement Sequelize Transactions, ensuring atomic database operations. "
        "Furthermore, due to the severe time constraints and varying availability of the 4 team members, scheduling synchronous meetings was difficult. "
        "We resolved this by adopting a strict asynchronous communication policy and utilizing modular code architecture so members could work independently without blocking one another."
    )

    doc.add_heading('3.5 Proposed Algorithms & Technologies Used', level=2)
    doc.add_paragraph("Algorithm 1: Escrow State Machine Algorithm", style='List Number')
    doc.add_paragraph(
        "When a booking is initiated, the algorithm verifies the client's wallet balance. If sufficient, funds are deducted and a 'ESCROW_HOLD' transaction is logged. "
        "Upon the Provider marking the job as 'COMPLETED', the algorithm executes a secure transfer, logging a 'PAYMENT_RECEIVED' transaction to the provider's wallet."
    )
    doc.add_paragraph("Technologies Used:")
    doc.add_paragraph("- Frontend: React, Tailwind CSS, Vite, Lucide React")
    doc.add_paragraph("- Backend: Node.js, Express.js, Socket.io")
    doc.add_paragraph("- Database: MySQL, Sequelize ORM")
    doc.add_paragraph("- Security: JSON Web Tokens (JWT), bcryptjs")
    doc.add_page_break()

    # RESULTS AND DISCUSSIONS
    doc.add_heading('Results and Discussions', level=1)
    
    doc.add_heading('Application Scenarios (Screenshots)', level=2)
    
    p1 = doc.add_paragraph()
    p1.add_run("[ ----------------------------------------------------- ]\n").bold = True
    p1.add_run("[ INSERT SCREENSHOT OF CLIENT DASHBOARD / SEARCH HERE ]\n").bold = True
    p1.add_run("[ ----------------------------------------------------- ]\n").bold = True
    doc.add_paragraph("Figure 4.1: The Client Dashboard showing the responsive grid of available service providers and category filtering.")

    p2 = doc.add_paragraph()
    p2.add_run("[ ----------------------------------------------------- ]\n").bold = True
    p2.add_run("[ INSERT SCREENSHOT OF WALLET TOP-UP MODAL HERE ]\n").bold = True
    p2.add_run("[ ----------------------------------------------------- ]\n").bold = True
    doc.add_paragraph("Figure 4.2: Mobile-responsive Wallet Modal simulating Mobile Money top-up with 9-digit phone validation.")

    p3 = doc.add_paragraph()
    p3.add_run("[ ----------------------------------------------------- ]\n").bold = True
    p3.add_run("[ INSERT SCREENSHOT OF REAL-TIME CHAT HERE ]\n").bold = True
    p3.add_run("[ ----------------------------------------------------- ]\n").bold = True
    doc.add_paragraph("Figure 4.3: The real-time chat interface connecting the client and provider via WebSockets.")

    doc.add_heading('API Request/Response Outputs', level=2)
    doc.add_paragraph("Below is an example of the JSON response generated by the GET /api/providers endpoint:")
    doc.add_paragraph(
        "{\n"
        '  "success": true,\n'
        '  "providers": [\n'
        '    {\n'
        '      "id": 1,\n'
        '      "user_id": 2,\n'
        '      "category": "Plumbing",\n'
        '      "hourly_rate": 5000,\n'
        '      "is_verified": true,\n'
        '      "user": {\n'
        '        "name": "Mark Doe",\n'
        '        "email": "mark@example.com"\n'
        '      }\n'
        '    }\n'
        '  ]\n'
        "}"
    )
    doc.add_page_break()

    # CHAPTER FOUR
    doc.add_heading('Chapter Four: Recommendations and Conclusion', level=1)
    
    doc.add_paragraph(
        "In conclusion, the 'Hire Me' project successfully achieved its core objective of delivering a secure, responsive, "
        "and reliable local service marketplace. By implementing a robust Escrow wallet system, real-time communication via WebSockets, "
        "and a rigorous provider verification matrix, the platform directly addresses the trust deficit inherent in the local gig economy. "
        "The cross-functional efforts of the four-member team allowed for the rapid deployment of these complex features utilizing the Scrum framework, "
        "resulting in a production-ready Minimum Viable Product."
    )
    
    doc.add_paragraph(
        "Despite the successful deployment, the team encountered notable difficulties throughout the software development lifecycle. "
        "The stringent time constraints placed immense pressure on the development cycles, requiring high efficiency and asynchronous collaboration. "
        "Technically, the migration from a local SQLite database to a robust MySQL environment introduced schema synchronization challenges. "
        "Furthermore, ensuring atomic database transactions during the Escrow payment flow required steep learning curves in ORM transaction management to prevent race conditions."
    )
    
    doc.add_paragraph(
        "For future iterations and further studies, it is highly recommended to integrate actual Mobile Money APIs (such as MTN or Orange Money) "
        "to replace the current simulated top-up system, thereby facilitating real-world financial operations. Additionally, incorporating an "
        "Artificial Intelligence-based matching algorithm could significantly enhance user experience by automatically recommending the most suitable "
        "service providers based on geographical proximity, past ratings, and budget constraints. Expanding the administrative dashboard to include "
        "predictive financial analytics would also provide greater oversight into platform economics."
    )

    doc.save('Hire_Me_Comprehensive_Project_Report.docx')
    print("Report generated successfully: Hire_Me_Comprehensive_Project_Report.docx")

if __name__ == '__main__':
    create_report()
