import docx
import os
import shutil

# 1. Modify Hire_Me_Comprehensive_Project_Report.docx
report_path = "Hire_Me_Comprehensive_Project_Report.docx"
doc = docx.Document(report_path)

# Add Team Roster
doc.add_heading('Project Team Members', level=1)
doc.add_paragraph('MEMBER 1: EBUA KEDZE FRANCK JORDAN', style='List Bullet')
doc.add_paragraph('Role: Scrum Master & Backend Lead')
doc.add_paragraph('MEMBER 2: AWONO FABIEN NYADINGA', style='List Bullet')
doc.add_paragraph('Role: Frontend Lead & DevOps Engineer')
doc.add_paragraph('MEMBER 3: Besomba Baonerges', style='List Bullet')
doc.add_paragraph('Role: Systems Analyst & Lead Technical Writer')
doc.add_paragraph('MEMBER 4: NDZI LEVI KONGNYU', style='List Bullet')
doc.add_paragraph('Role: UML Architect & Quality Assurance (QA) Engineer')

# Chapter 1
doc.add_heading('Chapter 1: Introduction', level=1)
doc.add_heading('1.1 General Introduction', level=2)
doc.add_paragraph('The rapid growth of the gig economy and the increasing reliance on digital marketplaces have transformed the way individuals and businesses connect for freelance services. In many emerging markets and local communities, finding reliable, skilled professionals for everyday tasks remains a highly fragmented and inefficient process. The "Hire Me" platform was conceived to bridge this gap by providing a secure, transparent, and user-friendly digital marketplace that seamlessly connects clients with verified local service providers.')

doc.add_heading('1.2 Problem Statement', level=2)
doc.add_paragraph('Currently, individuals seeking local services often rely on word-of-mouth recommendations or unverified online directories. This informal approach presents several significant challenges:')
doc.add_paragraph('1. Lack of Trust and Verification: Clients have no reliable way to verify the credentials, past performance, or criminal background of service providers.', style='List Bullet')
doc.add_paragraph('2. Payment Insecurity: Cash transactions and direct bank transfers leave both parties vulnerable to fraud, with clients risking incomplete work and providers risking non-payment.', style='List Bullet')
doc.add_paragraph('3. Inefficient Communication: Negotiating terms, discussing job requirements, and tracking service schedules is often scattered across various informal messaging apps.', style='List Bullet')
doc.add_paragraph('4. Poor Visibility for Professionals: Skilled workers struggle to market themselves effectively and build a verifiable digital reputation that reflects their expertise.', style='List Bullet')

doc.add_heading('1.3 Aim and Objectives', level=2)
doc.add_paragraph('The primary aim of the "Hire Me" project is to design, develop, and deploy a robust web-based service marketplace that facilitates secure and efficient transactions between clients and freelance professionals.')
doc.add_paragraph('Specific Objectives:')
doc.add_paragraph('Secure Financial Transactions: To implement a robust wallet and escrow system that holds client funds securely until the agreed-upon job is completed satisfactorily, ensuring fairness for both parties.', style='List Bullet')
doc.add_paragraph('Identity and Skill Verification: To provide an administrative backend where provider credentials and identification documents can be reviewed and officially verified.', style='List Bullet')
doc.add_paragraph('Real-Time Communication: To integrate a real-time messaging system allowing clients and providers to discuss job specifications without exchanging personal phone numbers.', style='List Bullet')
doc.add_paragraph('Reputation Management: To build a transparent review and rating system that holds providers accountable and helps clients make informed decisions based on past performance.', style='List Bullet')
doc.add_paragraph('Role-Based Access Control: To create distinct, tailored dashboard experiences for Clients, Providers, and Administrators, ensuring that each user type has access to the specific tools they need.', style='List Bullet')

# Chapter 2
doc.add_heading('Chapter 2: Literature Review and Methodology', level=1)
doc.add_heading('2.1 Review of Related Literature', level=2)
doc.add_paragraph('The evolution of freelance service marketplaces has been well-documented in recent software engineering and economic literature. Platforms like Upwork, Fiverr, and TaskRabbit have demonstrated the viability of connecting independent contractors with clients via centralized digital hubs. However, literature highlights several persistent pain points in these platforms, particularly regarding trust, quality assurance, and high commission fees.')
doc.add_paragraph('Decentralized trust mechanisms—such as escrow services—are critical in mitigating fraud. The "Hire Me" platform addresses this directly by integrating a built-in virtual wallet and an escrow holding system, ensuring that funds are committed before work begins but not released until client satisfaction is confirmed. Furthermore, localized platforms face unique challenges related to real-time coordination, which "Hire Me" solves through integrated WebSocket-based chat and instant status updates, reducing the friction often cited in user experience studies of service marketplaces.')

doc.add_heading('2.2 Software Development Methodologies', level=2)
doc.add_paragraph('In software engineering, the choice of development methodology significantly impacts the project\'s success, adaptability, and time-to-market. The two primary paradigms considered for this project were the Waterfall model and Agile methodologies.')
doc.add_paragraph('Waterfall Methodology: A linear, sequential approach where each phase (Requirements, Design, Implementation, Testing, Deployment) must be completed before the next begins. While offering strict documentation and predictability, it lacks flexibility. If a critical flaw in the escrow logic was discovered during testing, revisiting the design phase would be highly disruptive.', style='List Bullet')
doc.add_paragraph('Agile Methodology: An iterative approach that promotes continuous delivery, flexible responses to change, and collaborative development.', style='List Bullet')

doc.add_heading('2.3 Justification for Choosing Scrum', level=2)
doc.add_paragraph('For the development of the "Hire Me" platform, the Scrum framework—a specific subset of Agile—was selected. This decision was driven by several factors:')
doc.add_paragraph('Iterative Feature Delivery: The platform required complex, interwoven features (e.g., authentication, wallet, real-time chat). Scrum allowed the team to break these down into manageable one-week Sprints, delivering a functional increment before adding complexities.', style='List Number')
doc.add_paragraph('Rapid Adaptation: As the team encountered technical challenges, Scrum\'s daily stand-ups and Sprint Retrospectives allowed for immediate pivoting and architecture adjustments without derailing the entire project timeline.', style='List Number')
doc.add_paragraph('Clear Role Distribution: With a clear division of labor across frontend, backend, and documentation, the Scrum framework provided necessary structure. The Product Backlog ensured that all contributors remained aligned on priority tasks, preventing scope creep.', style='List Number')

# Chapter 3
doc.add_heading('Chapter 3: System Design and Implementation', level=1)
doc.add_paragraph('This chapter presents the system design through UML diagrams and showcases the implementation through screenshots of the final software.')
doc.add_heading('3.1 UML Diagrams', level=2)
doc.add_paragraph('[INSERT USE CASE DIAGRAM HERE: Lead Architect - NDZI LEVI KONGNYU]')
doc.add_paragraph('[INSERT CLASS DIAGRAM HERE: Lead Architect - NDZI LEVI KONGNYU]')
doc.add_paragraph('[INSERT SEQUENCE DIAGRAM HERE: Lead Architect - NDZI LEVI KONGNYU]')
doc.add_paragraph('[INSERT ENTITY-RELATIONSHIP DIAGRAM (ERD) HERE: Lead Architect - NDZI LEVI KONGNYU]')
doc.add_heading('3.2 Software Screenshots', level=2)
doc.add_paragraph('[INSERT SCREENSHOT: LANDING PAGE AND AUTHENTICATION]')
doc.add_paragraph('[INSERT SCREENSHOT: CLIENT DASHBOARD]')
doc.add_paragraph('[INSERT SCREENSHOT: PROVIDER PROFILE AND BOOKING MODAL]')
doc.add_paragraph('[INSERT SCREENSHOT: REAL-TIME CHAT INTERFACE]')
doc.add_paragraph('[INSERT SCREENSHOT: PROVIDER DASHBOARD (EARNINGS & BOOKINGS)]')
doc.add_paragraph('[INSERT SCREENSHOT: ADMIN DASHBOARD (ANALYTICS & MODERATION)]')

# Chapter 4
doc.add_heading('Chapter 4: Recommendations and Conclusion', level=1)
doc.add_heading('4.1 Summary of Achievements', level=2)
doc.add_paragraph('The "Hire Me" platform successfully achieved its primary objective of creating a secure, intuitive, and feature-rich digital marketplace for freelance professionals and clients. Over the course of four development Sprints, the team successfully engineered a full-stack application utilizing Node.js, Express, Sequelize, and React. Core functionalities—including role-based access control, a secure virtual wallet with escrow capabilities, an interactive booking lifecycle, and real-time Socket.io chat—were fully implemented and tested. The resulting system effectively addresses the lack of trust and communication inefficiencies prevalent in informal local service sectors by providing transparent reviews, verified badges, and structured financial holding mechanisms.')

doc.add_heading('4.2 Difficulties Encountered', level=2)
doc.add_paragraph('Despite the project\'s success, the team faced several technical challenges during development. Managing the complex state of the virtual wallet during edge cases—such as booking cancellations and dispute resolutions—required extensive transaction rollback testing to ensure financial data integrity. Additionally, integrating real-time WebSockets with the React frontend presented challenges regarding state synchronization; specifically, ensuring that chat messages appeared optimistically for the sender while preventing duplicate broadcasts from the server. Finally, coordinating code merges across frontend and backend branches required strict version control discipline to prevent configuration conflicts.')

doc.add_heading('4.3 Future Recommendations', level=2)
doc.add_paragraph('While the current iteration of the platform serves as a robust Minimum Viable Product (MVP), several features are recommended for future development to enhance scalability and user experience. First, integrating a third-party payment gateway (such as Stripe or Mobile Money APIs) would allow users to fund their virtual wallets directly using real-world currency. Second, implementing geographical map views (using Google Maps API) would allow clients to visualize the proximity of nearby providers. Finally, the introduction of a monthly calendar view for providers to manage their availability, alongside automated PDF invoice generation for completed jobs, would significantly elevate the platform\'s utility for full-time freelance professionals.')

doc.save(report_path)

# Move Report to documentation folder
if not os.path.exists('documentation'):
    os.makedirs('documentation')

shutil.move(report_path, os.path.join('documentation', report_path))


# 2. Generate Product Backlog
pb_doc = docx.Document()
pb_doc.add_heading('Hire Me Platform - Product Backlog', level=1)
pb_doc.add_paragraph('The Product Backlog is a comprehensive and prioritized list of all features, functionalities, and requirements necessary for the Hire Me platform. It serves as the single source of truth for the development team.')

pb_doc.add_heading('High Priority (MVP Features)', level=2)
table = pb_doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'Feature'
hdr_cells[2].text = 'User Story'
hdr_cells[3].text = 'Priority'
hdr_cells[4].text = 'Status'

records = [
    ('PB-01', 'User Authentication', 'As a user, I want to register and log in securely so that my personal data is protected.', 'High', 'Done'),
    ('PB-02', 'Role-Based Dashboards', 'As a user, I want to see a tailored dashboard depending on whether I am a Client, Provider, or Admin.', 'High', 'Done'),
    ('PB-03', 'Provider Profiles', 'As a professional, I want to set up my profile (trade, hourly rate, bio, location) so clients can hire me.', 'High', 'Done'),
    ('PB-04', 'Wallet & Escrow System', 'As a client, I want to fund my wallet and have funds held in escrow until a job is completed securely.', 'High', 'Done'),
    ('PB-05', 'Booking System', 'As a client, I want to request a service provider for a specific date and time.', 'High', 'Done'),
    ('PB-06', 'Booking Management', 'As a provider, I want to accept or reject incoming booking requests.', 'High', 'Done'),
    ('PB-07', 'Job Completion & Payout', 'As a provider, I want to receive my funds directly to my wallet (minus a 5% platform fee) when a client marks a job as completed.', 'High', 'Done'),
    ('PB-08', 'Review & Rating System', 'As a client, I want to leave a rating and review for a provider after a completed job.', 'High', 'Done'),
]
for id, feature, story, priority, status in records:
    row_cells = table.add_row().cells
    row_cells[0].text = id
    row_cells[1].text = feature
    row_cells[2].text = story
    row_cells[3].text = priority
    row_cells[4].text = status

pb_doc.add_heading('Medium Priority (Enhancements)', level=2)
table2 = pb_doc.add_table(rows=1, cols=5)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'ID'
hdr_cells2[1].text = 'Feature'
hdr_cells2[2].text = 'User Story'
hdr_cells2[3].text = 'Priority'
hdr_cells2[4].text = 'Status'
records2 = [
    ('PB-09', 'Real-Time Chat', 'As a user, I want to chat with my matched client/provider in real-time to discuss job details.', 'Medium', 'Done'),
    ('PB-10', 'Save/Favorite Providers', 'As a client, I want to save my favorite providers so I can quickly book them again in the future.', 'Medium', 'Done'),
    ('PB-11', 'Provider Verification', 'As an admin, I want to verify professional credentials and assign a "Verified" badge.', 'Medium', 'Done'),
    ('PB-12', 'Earnings Dashboard', 'As a provider, I want to view my past earnings, pending escrow funds, and transaction history.', 'Medium', 'Done'),
    ('PB-13', 'Admin Analytics', 'As an admin, I want to view platform statistics, user growth, and revenue from platform fees.', 'Medium', 'Done'),
    ('PB-14', 'Booking Cancellation', 'As a client, I want to cancel a pending booking and receive an instant refund from escrow.', 'Medium', 'Done'),
]
for id, feature, story, priority, status in records2:
    row_cells = table2.add_row().cells
    row_cells[0].text = id
    row_cells[1].text = feature
    row_cells[2].text = story
    row_cells[3].text = priority
    row_cells[4].text = status

pb_path = os.path.join('documentation', 'Product_Backlog.docx')
pb_doc.save(pb_path)


# 3. Generate Sprint Backlogs
sb_doc = docx.Document()
sb_doc.add_heading('Hire Me Platform - Sprint Backlogs', level=1)
sb_doc.add_paragraph('The development of the Hire Me platform was divided into a 4-week timeline. We utilized the Scrum methodology to iteratively build, test, and deploy features. Below is the breakdown of tasks across four one-week Sprints.')

sprints = [
    ('Sprint 1: Foundation and Architecture', 'Establish the project architecture, database models, and core authentication system.', [
        ('Initialize backend (Node.js/Express) and frontend (React/Vite) repositories.', 'AWONO FABIEN NYADINGA, EBUA KEDZE FRANCK JORDAN', 'Done'),
        ('Design database schema (Users, Profiles, Bookings, Transactions, Messages).', 'EBUA KEDZE FRANCK JORDAN', 'Done'),
        ('Implement JWT Authentication and Role-Based Access Control (RBAC).', 'EBUA KEDZE FRANCK JORDAN', 'Done'),
        ('Create basic UI wireframes and establish the Tailwind CSS design system.', 'AWONO FABIEN NYADINGA', 'Done'),
        ('Build Landing Auth page and basic routing structure.', 'AWONO FABIEN NYADINGA', 'Done'),
    ]),
    ('Sprint 2: Core Marketplace Workflows', 'Implement the primary flows: Profiles, Wallets, and the Booking Lifecycle.', [
        ('Develop the Wallet & Escrow system endpoints (Top-up, Escrow Hold, Payout).', 'EBUA KEDZE FRANCK JORDAN', 'Done'),
        ('Build Provider Search API with filtering by category and hourly rate.', 'EBUA KEDZE FRANCK JORDAN', 'Done'),
        ('Develop Client Dashboard (Provider search, listing bookings).', 'AWONO FABIEN NYADINGA', 'Done'),
        ('Develop Provider Dashboard (Accept/Reject bookings, Earnings view).', 'AWONO FABIEN NYADINGA', 'Done'),
        ('Implement Provider Profile page and Booking Modal.', 'AWONO FABIEN NYADINGA', 'Done'),
    ]),
    ('Sprint 3: Enhancements and Real-Time Features', 'Add real-time communication, reviews, and administrative oversight.', [
        ('Integrate Socket.io for real-time notifications (booking_updated).', 'EBUA KEDZE FRANCK JORDAN', 'Done'),
        ('Develop Real-Time Chat Modal with optimistic UI updates.', 'AWONO FABIEN NYADINGA', 'Done'),
        ('Implement Review and Rating logic (updates provider average rating).', 'EBUA KEDZE FRANCK JORDAN', 'Done'),
        ('Build Admin Dashboard (User moderation, analytics, verifying credentials).', 'AWONO FABIEN NYADINGA', 'Done'),
        ('Draft Project Report Chapter 1 & 2.', 'Besomba Baonerges', 'Done'),
    ]),
    ('Sprint 4: Polish, UI Consistency, and Documentation', 'Finalize UI features, fix bugs, and complete academic documentation.', [
        ('Implement "Save/Favorite Provider" feature with optimistic UI.', 'AWONO FABIEN NYADINGA, EBUA KEDZE FRANCK JORDAN', 'Done'),
        ('Fix Socket.io duplicate message bugs and cancel booking edge cases.', 'EBUA KEDZE FRANCK JORDAN', 'Done'),
        ('Finalize UML diagrams (Class, Sequence, Use Case, ERD).', 'NDZI LEVI KONGNYU', 'Done'),
        ('Finalize Project Management artifacts (Backlogs) and remaining report chapters.', 'Besomba Baonerges', 'Done'),
    ])
]

for title, goal, tasks in sprints:
    sb_doc.add_heading(title, level=2)
    sb_doc.add_paragraph(f'Goal: {goal}')
    table = sb_doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Task'
    hdr[1].text = 'Assignee'
    hdr[2].text = 'Status'
    for t, a, s in tasks:
        row = table.add_row().cells
        row[0].text = t
        row[1].text = a
        row[2].text = s

sb_path = os.path.join('documentation', 'Sprint_Backlogs.docx')
sb_doc.save(sb_path)

print("DOCX generation complete.")
