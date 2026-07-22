from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_api_doc():
    doc = Document()
    
    # Title Page
    doc.add_heading('AUTOMATED API DOCUMENTATION', 0)
    title = doc.add_paragraph('Project Title: Hire Me - A Local Service Marketplace\n')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Introduction
    doc.add_heading('1. Swagger Automation Overview', level=1)
    doc.add_paragraph(
        "As per the Phase 4 requirements, the 'Hire Me' platform integrates automated API documentation using the Swagger (OpenAPI 3.0) framework. "
        "The documentation is generated dynamically at runtime using the 'swagger-jsdoc' and 'swagger-ui-express' packages in Node.js. "
        "When the server is running, the interactive UI is hosted at http://localhost:5000/api-docs."
    )
    doc.add_paragraph("This document serves as a static snapshot of the automated endpoints available in the system.")

    # Endpoints
    doc.add_heading('2. Core API Endpoints', level=1)

    # Function to create endpoint block
    def add_endpoint(doc, method, url, desc, auth, params, body, response):
        p = doc.add_paragraph()
        p.add_run(f"[{method.upper()}] ").bold = True
        p.add_run(url).bold = True
        
        doc.add_paragraph(f"Description: {desc}")
        doc.add_paragraph(f"Requires Authentication: {auth}")
        
        if params:
            doc.add_paragraph("Query Parameters: " + params)
        if body:
            doc.add_paragraph("Request Body (JSON):")
            doc.add_paragraph(body, style='Intense Quote')
            
        doc.add_paragraph("Expected Response:")
        doc.add_paragraph(response, style='Intense Quote')
        doc.add_paragraph("-" * 50)

    add_endpoint(
        doc, "POST", "/api/auth/register",
        "Registers a new user (Client or Provider) and returns a JWT token.",
        "No", None,
        '{\n  "name": "John Doe",\n  "email": "john@example.com",\n  "password": "secret123",\n  "role": "CLIENT"\n}',
        'HTTP 201 Created\n{\n  "token": "eyJhbG...",\n  "user": { ... }\n}'
    )

    add_endpoint(
        doc, "POST", "/api/auth/login",
        "Authenticates a user and returns a JWT token.",
        "No", None,
        '{\n  "email": "john@example.com",\n  "password": "secret123"\n}',
        'HTTP 200 OK\n{\n  "token": "eyJhbG...",\n  "user": { ... }\n}'
    )

    add_endpoint(
        doc, "GET", "/api/providers",
        "Fetches a list of verified service providers.",
        "No", "?category=Plumbing&maxPrice=10000",
        None,
        'HTTP 200 OK\n{\n  "success": true,\n  "providers": [ ... ]\n}'
    )

    add_endpoint(
        doc, "POST", "/api/wallet/topup",
        "Simulates a Mobile Money transaction to top-up a user's wallet.",
        "Yes (JWT)", None,
        '{\n  "amount": 10000,\n  "provider": "MTN"\n}',
        'HTTP 200 OK\n{\n  "success": true,\n  "new_balance": 10000\n}'
    )

    add_endpoint(
        doc, "POST", "/api/bookings",
        "Creates a new booking, atomically deducting funds from the client into Escrow.",
        "Yes (JWT)", None,
        '{\n  "provider_id": 2,\n  "amount": 5000,\n  "job_date": "2026-10-15"\n}',
        'HTTP 201 Created\n{\n  "success": true,\n  "booking": { ... }\n}'
    )

    add_endpoint(
        doc, "PATCH", "/api/bookings/{id}/status",
        "Updates booking status (e.g., to COMPLETED) and automatically releases Escrow funds to the provider.",
        "Yes (JWT)", None,
        '{\n  "status": "COMPLETED"\n}',
        'HTTP 200 OK\n{\n  "success": true,\n  "message": "Status updated successfully"\n}'
    )

    doc.add_page_break()
    doc.add_heading('3. Screenshots of Automated Swagger UI', level=1)
    
    p1 = doc.add_paragraph()
    p1.add_run("[ ----------------------------------------------------- ]\n").bold = True
    p1.add_run("[ INSERT SCREENSHOT OF http://localhost:5000/api-docs ]\n").bold = True
    p1.add_run("[ ----------------------------------------------------- ]\n").bold = True
    doc.add_paragraph("Figure 1.1: The interactive Swagger UI dashboard showing all available endpoints.")

    doc.save('Hire_Me_Automated_API_Documentation.docx')
    print("API Document generated successfully: Hire_Me_Automated_API_Documentation.docx")

if __name__ == '__main__':
    create_api_doc()
