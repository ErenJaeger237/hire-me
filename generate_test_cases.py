from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_cases_document():
    doc = Document()
    
    # Title Page
    doc.add_heading('SYSTEM TEST CASE DOCUMENT', 0)
    title = doc.add_paragraph('Project Title: Hire Me - A Local Service Marketplace\n')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Introduction
    doc.add_heading('1. Test Objectives', level=1)
    doc.add_paragraph(
        "The objective of this document is to outline the formal test cases utilized to validate the functional "
        "requirements of the 'Hire Me' platform. These test cases ensure that core features such as Authentication, "
        "Role-Based Access Control, Escrow Booking, and the Digital Wallet operate correctly under both expected "
        "and edge-case conditions."
    )

    doc.add_heading('2. Formal Test Cases', level=1)

    # Function to add a test case table
    def add_test_case(doc, tc_id, desc, pre, steps, expected, actual, status):
        doc.add_heading(f'Test Case: {tc_id} - {desc}', level=2)
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Write data
        table.cell(0, 0).text = 'Pre-conditions'
        table.cell(0, 1).text = pre
        
        table.cell(1, 0).text = 'Test Steps'
        table.cell(1, 1).text = steps
        
        table.cell(2, 0).text = 'Expected Result'
        table.cell(2, 1).text = expected
        
        table.cell(3, 0).text = 'Actual Result'
        table.cell(3, 1).text = actual
        
        table.cell(4, 0).text = 'Status'
        status_cell = table.cell(4, 1)
        status_cell.text = status
        
        doc.add_paragraph("\n")

    # TC 1
    add_test_case(
        doc, "TC-01", "User Authentication (Invalid Credentials)",
        "User is on the Login Page.",
        "1. Enter unregistered email.\n2. Enter random password.\n3. Click 'Login'.",
        "System displays 'Invalid credentials' error. Access is denied.",
        "System displays 'Invalid credentials' error. Access is denied.",
        "PASS"
    )

    # TC 2
    add_test_case(
        doc, "TC-02", "Role-Based Access Control (RBAC)",
        "User is logged in with the 'CLIENT' role.",
        "1. Navigate manually to '/admin-dashboard' URL.",
        "System intercepts the route, identifies invalid role, and redirects to Client Dashboard or shows 'Unauthorized'.",
        "System blocks access and shows standard interface.",
        "PASS"
    )

    # TC 3
    add_test_case(
        doc, "TC-03", "Wallet Top-Up (Validation)",
        "User is logged in and clicks 'Wallet'.",
        "1. Enter 5000 in Amount.\n2. Enter an 8-digit phone number (e.g., 67012345).\n3. Click 'Confirm Top Up'.",
        "System prevents submission and displays 'Phone number must be exactly 9 digits'.",
        "System prevents submission and displays validation error.",
        "PASS"
    )

    # TC 4
    add_test_case(
        doc, "TC-04", "Wallet Top-Up (Success)",
        "User is logged in and clicks 'Wallet'.",
        "1. Enter 10000 in Amount.\n2. Enter a valid 9-digit Cameroonian phone number.\n3. Click 'Confirm Top Up'.",
        "System processes request, creates a 'TOPUP' transaction, and increments user's wallet_balance by 10000.",
        "Wallet balance updates instantly to reflect the +10000 FCFA.",
        "PASS"
    )

    # TC 5
    add_test_case(
        doc, "TC-05", "Booking Escrow (Insufficient Funds)",
        "Client has 2000 FCFA in wallet. Provider rate is 5000 FCFA.",
        "1. Client clicks 'Book' on the Provider.\n2. Confirm the booking modal.",
        "System rejects the booking, displays 'Insufficient funds', and prompts user to top-up.",
        "System rejects the booking and displays 'Insufficient funds'.",
        "PASS"
    )

    # TC 6
    add_test_case(
        doc, "TC-06", "Booking Escrow (Sufficient Funds)",
        "Client has 10000 FCFA in wallet. Provider rate is 5000 FCFA.",
        "1. Client clicks 'Book' on the Provider.\n2. Confirm the booking modal.",
        "System creates Booking, deducts 5000 FCFA from Client, logs 'ESCROW_HOLD' transaction, and notifies Provider.",
        "Client balance drops to 5000 FCFA. Booking appears in Provider's Active Jobs.",
        "PASS"
    )

    # TC 7
    add_test_case(
        doc, "TC-07", "Job Completion & Escrow Release",
        "Booking is in 'PENDING'/'IN_PROGRESS' state. Funds are held in Escrow.",
        "1. Provider clicks 'Mark Job Completed'.",
        "System changes booking status to 'COMPLETED', creates 'PAYMENT_RECEIVED' transaction, and atomically adds 5000 FCFA to Provider's wallet.",
        "Provider balance increases by 5000 FCFA. Job moves to Completed Jobs list.",
        "PASS"
    )

    doc.save('Hire_Me_Test_Case_Document.docx')
    print("Test Case Document generated successfully: Hire_Me_Test_Case_Document.docx")

if __name__ == '__main__':
    create_test_cases_document()
